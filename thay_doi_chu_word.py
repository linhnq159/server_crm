from docx import Document


# for section in document.sections:
#     for paragraph in section.footer.paragraphs:
#         print(paragraph.text)

# for section in document.sections:
#     for table in section.footer.tables:
#         for row in table.rows:
#             for cell in row.cells:
#                 for paragraph in cell.paragraphs:
#                     print(paragraph.text)

# for paragraph in doc.paragraphs:
#     print(paragraph.text)
#     inline = paragraph.runs
#     for i in range(len(inline)):
#         print(inline[i].text)

# for table in doc.tables:
#     for row in table.rows:
#         for cell in row.cells:
#             for paragraph in cell.paragraphs:
#                 print(paragraph.text)
#                 if text_change in paragraph.text:
#                     inline = paragraph.runs
#                     for i in range(len(inline)):
#                         print(inline[i].text)

def thay_doi_chu_word(output_file ,text_change , list_new_text):
    doc = Document(output_file)
    p = -1
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if text_change in paragraph.text:
                        p += 1
                        # print(paragraph.text)
                        new_text = list_new_text[p]
                        if new_text != ' ':
                            # Tìm vị trí từ cần thay đổi
                            inline = paragraph.runs
                            text_change_split = text_change.split()
                            len_text_change = len(text_change_split)
                            in_line = []
                            list_change = []
                            j = 0
                            # Loop added to work with runs (strings with same style)
                            for i in range(len(inline)):
                                while j < len_text_change and inline[i].text.strip() != '':
                                    if text_change_split[j] in inline[i].text:
                                        if j+1 == len_text_change:
                                            list_change.append(j)
                                            in_line.append(i)
                                            break
                                        list_change.append(j)
                                        in_line.append(i)
                                        j += 1
                                        continue
                                    break

                            # Thay đổi chữ cần thay đổi
                            for i in range(1,len(in_line)):
                                inline[in_line[i]].text = inline[in_line[i]].text.replace(text_change_split[i], "")
                            inline[in_line[0]].text = inline[in_line[0]].text.replace(text_change_split[0], new_text)


    for paragraph in doc.paragraphs:
        if text_change in paragraph.text:
            p += 1
            # print(paragraph.text)
            new_text = list_new_text[p]
            if new_text != " ":
                inline = paragraph.runs
                text_change_split = text_change.split()
                len_text_change = len(text_change_split)
                in_line = []
                list_change = []
                j = 0
                # Loop added to work with runs (strings with same style)
                for i in range(len(inline)):
                    while j < len_text_change and inline[i].text.strip() != '':
                        if text_change_split[j] in inline[i].text:
                            if j + 1 == len_text_change:
                                list_change.append(j)
                                in_line.append(i)
                                break
                            list_change.append(j)
                            in_line.append(i)
                            j += 1
                            continue
                        break

                for i in range(1, len(in_line)):
                    inline[in_line[i]].text = inline[in_line[i]].text.replace(text_change_split[i], "")
                inline[in_line[0]].text = inline[in_line[0]].text.replace(text_change_split[0], new_text)

    doc.save(output_file)

